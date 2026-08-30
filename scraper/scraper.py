"""
scraper.py — fetch jobs from LinkedIn via the harvestapi/linkedin-job-search Apify actor.

run() reads cfg["roles"] (set by the user in the dashboard UI and saved to
config.yaml) to determine what to search for.

extract_roles_from_resume() is a separate helper used only by the /api/roles
dashboard endpoint to suggest an initial role list from the resume PDF — it
is never called during the scraping pipeline.

Each call is filtered to jobs posted in the last 24 hours and the total is
capped at _MAX_TOTAL_JOBS across all role/location combinations.

Actor ID is read from env var APIFY_LINKEDIN_ACTOR so it can be swapped
without touching code.
"""

import json
import logging
import os
import re
import sqlite3
import time
from itertools import product
from pathlib import Path

import anthropic
import pdfplumber
from apify_client import ApifyClient
from dotenv import load_dotenv

from core.config import load_config

load_dotenv()
log = logging.getLogger(__name__)


def _safe_str(val) -> str:
    """Convert any API field value to a plain string, handling nested dicts."""
    if val is None:
        return ""
    if isinstance(val, dict):
        return str(val.get("name") or val.get("value") or val.get("text") or "")
    return str(val)


def _extract_company(raw) -> str:
    if isinstance(raw, dict):
        return raw.get("name") or raw.get("companyName") or ""
    return str(raw) if raw else ""

_LINKEDIN_ACTOR  = os.getenv("APIFY_LINKEDIN_ACTOR", "harvestapi/linkedin-job-search")
_MODEL           = "claude-sonnet-4-6"

_MAX_TOTAL_JOBS = 100  # hard cap on total jobs returned per daily run

_scrape_progress: dict = {
    "total_combinations": 0,
    "completed_combinations": 0,
    "current_search": "",
    "jobs_found_so_far": 0,
}

# Stop signal handling (file-based for inter-process communication)
_STOP_FLAG_FILE = Path(__file__).parent.parent / "uploads" / ".stop_scraping"
_current_apify_run_id = None


def get_scrape_progress() -> dict:
    return dict(_scrape_progress)


def _should_stop_scraping() -> bool:
    """Check if stop flag file exists (can be called from subprocess)."""
    return _STOP_FLAG_FILE.exists()


def stop_scraping() -> None:
    """Signal the scraper to stop and kill any running Apify actor."""
    global _current_apify_run_id

    # Create stop flag file so subprocess detects it
    _STOP_FLAG_FILE.parent.mkdir(parents=True, exist_ok=True)
    _STOP_FLAG_FILE.touch()
    log.info("⏹️ Stop signal created at %s", _STOP_FLAG_FILE)

    if _current_apify_run_id:
        try:
            client = _client()
            log.info("Aborting Apify actor run: %s", _current_apify_run_id)
            client.run(_current_apify_run_id).abort()
            log.info("✓ Apify actor %s aborted successfully", _current_apify_run_id)
        except Exception as exc:
            log.warning("Failed to abort Apify actor %s: %s", _current_apify_run_id, exc)
        finally:
            _current_apify_run_id = None
    else:
        log.info("No active Apify actor to abort")


def reset_stop_flag() -> None:
    """Reset the stop flag for a new scraping session."""
    try:
        if _STOP_FLAG_FILE.exists():
            _STOP_FLAG_FILE.unlink()
            log.debug("Cleared stop flag file")
    except Exception:
        pass

# Titles containing these compounds are kept regardless of other keywords.
# Rationale: "Java/Python Backend Engineer" lists Python as secondary — still relevant.
_SAFE_TITLE_COMPOUNDS = [
    "backend engineer", "backend developer",
    "software engineer", "software developer",
    "full stack", "fullstack", "full-stack",
    "platform engineer",
]

_SKIP_TITLE_KEYWORDS = [
    # Wrong primary-stack roles (Java candidate is unlikely to match these)
    "python developer", "python engineer",
    "go developer", "golang",
    "ruby developer", "ruby on rails",
    "php developer", "php engineer",
    "frontend developer", "frontend engineer",
    "react developer", "angular developer", "vue developer",
    "ios developer", "android developer",
    "mobile developer", "mobile engineer",
    # Non-tech roles
    "marketing", "sales", "accountant", "mechanic", "nurse", "driver",
    "designer", "hr ", " hr", "recruiter", "lawyer", "finance", "intern",
    "co-founder", "assistant", "coordinator", "consultant", "analyst",
    "manager", "director", "product owner", "scrum master",
]

_TECH_TITLE_KEYWORDS = [
    "engineer", "developer", "architect", "backend", "software", "java",
    "python", "devops", "cloud", "data", "fullstack", "full-stack", "api",
    "platform", "infrastructure", "site reliability", "sre", "ml ",
    "machine learning", "kotlin", "scala", "typescript",
]


def is_title_relevant(job_title: str, search_role: str) -> bool:
    """Fast keyword check — no API calls. Returns False for clearly off-topic titles."""
    title = job_title.lower()
    # Safe compounds bypass the skip list (e.g. "Python Backend Engineer" → keep)
    if any(kw in title for kw in _SAFE_TITLE_COMPOUNDS):
        return True
    if any(kw in title for kw in _SKIP_TITLE_KEYWORDS):
        return False
    if any(kw in title for kw in _TECH_TITLE_KEYWORDS):
        return True
    # Fall back to checking overlap with the role's own words
    role_words = set(search_role.lower().split())
    return bool(role_words & set(title.split()))

_extracted_roles_cache: list[str] | None = None

_ROLE_SYSTEM = """\
You analyse a candidate's resume and extract job-search terms.
Return ONLY a JSON array of strings — no prose, no markdown fences.
"""

_ROLE_PROMPT = """\
## Resume
{resume}

Extract up to 6 job titles or core skill areas that best represent this candidate for job searching.
Rules:
- Only include titles/skills that are clearly evidenced in the resume.
- Do NOT invent roles absent from the resume.
- Use standard industry terms a recruiter would type into a job board.
- Prefer specific titles over generic ones (e.g. "Data Engineer" over "Engineer").

Return a JSON array, e.g.: ["Data Engineer", "Python Developer", "Backend Engineer"]
"""


def extract_roles_from_resume(cfg: dict) -> list[str]:
    """Read the resume (PDF or DOCX) and ask Claude for up to 6 job-search terms."""
    global _extracted_roles_cache
    if _extracted_roles_cache is not None:
        return _extracted_roles_cache

    resume_path = cfg["paths"].get("resume_en_docx") or cfg["paths"]["resume_en"]
    if not resume_path.exists():
        resume_path = cfg["paths"]["resume_en"]
    if not resume_path.exists():
        raise FileNotFoundError(f"Resume not found: {resume_path}")

    if resume_path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(resume_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        resume_text = "\n".join(parts)
    else:
        text_parts: list[str] = []
        with pdfplumber.open(resume_path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        resume_text = "\n".join(text_parts)

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise EnvironmentError("ANTHROPIC_API_KEY is not set in .env")

    client = anthropic.Anthropic(api_key=api_key)
    resp = client.messages.create(
        model=_MODEL,
        max_tokens=256,
        system=_ROLE_SYSTEM,
        messages=[{"role": "user", "content": _ROLE_PROMPT.format(resume=resume_text[:4000])}],
    )
    raw = resp.content[0].text.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)
        raw = raw.strip()
    try:
        roles = json.loads(raw)
        if not isinstance(roles, list):
            raise ValueError("Expected JSON array")
    except Exception:
        log.warning("Role extraction parse error — raw: %s", raw[:200])
        roles = ["Data Engineer", "Python Developer"]

    roles = [r.strip() for r in roles if isinstance(r, str) and r.strip()][:6]
    log.info("Extracted roles from resume: %s", roles)
    _extracted_roles_cache = roles
    return roles


def _client() -> ApifyClient:
    token = os.getenv("APIFY_API_TOKEN")
    if not token:
        raise EnvironmentError("APIFY_API_TOKEN is not set in .env")
    return ApifyClient(token)


def _scrape_linkedin(client: ApifyClient, role: str, location: str,
                     cfg: dict | None = None) -> list[dict]:
    global _current_apify_run_id

    log.info("LinkedIn scrape: %s @ %s", role, location)
    _cfg = cfg or {}
    posted_limit = _cfg.get("posted_limit", "24h")
    count = int(_cfg.get("scrape_pool_size", 25))
    experience_levels = _cfg.get("experience_levels", ["mid-senior", "senior"])
    run_input = {
        "jobTitles":       [role],
        "locations":       [location],
        "maxItems":        count,
        "postedLimit":     posted_limit,
        "employmentType":  ["full-time"],
        "experienceLevel": experience_levels,
        "sortBy":          "relevance",
    }
    log.info("Apify run_input: %s", run_input)

    # Check stop flag before starting actor
    if _should_stop_scraping():
        log.info("Stop flag detected before starting actor — aborting this combo")
        raise RuntimeError("Scraping stopped by user")

    run = client.actor(_LINKEDIN_ACTOR).start(run_input=run_input)
    _current_apify_run_id = run.id
    log.info("[Apify] Run started: %s", _current_apify_run_id)

    # Poll until the actor finishes, streaming its logs to agent.log in real-time
    _last_log_chars = 0
    _poll_interval  = 5  # seconds between status checks
    while True:
        if _should_stop_scraping():
            log.info("[Apify] Stop flag — aborting actor %s", _current_apify_run_id)
            try:
                client.run(_current_apify_run_id).abort()
                log.info("[Apify] Aborted %s", _current_apify_run_id)
            except Exception as exc:
                log.warning("[Apify] Abort failed: %s", exc)
            _current_apify_run_id = None
            raise RuntimeError("Scraping stopped by user")

        # Fetch any new log lines from the Apify platform and write to agent.log
        try:
            log_text: str = client.run(_current_apify_run_id).log().get() or ""
            if len(log_text) > _last_log_chars:
                for _line in log_text[_last_log_chars:].splitlines():
                    if _line.strip():
                        log.info("[Apify] %s", _line.strip())
                _last_log_chars = len(log_text)
        except Exception:
            pass

        run_info = client.run(_current_apify_run_id).get()
        status   = run_info.status if run_info else "RUNNING"
        if status not in ("RUNNING", "READY"):
            log.info("[Apify] Run finished with status: %s", status)
            break

        time.sleep(_poll_interval)

    run = client.run(_current_apify_run_id).get() or run
    _current_apify_run_id = None

    raw_items: list[dict] = list(client.dataset(run.default_dataset_id).iterate_items())
    _current_apify_run_id = None  # Clear after actor finishes
    log.info("  → %d raw items from Apify dataset", len(raw_items))
    if raw_items:
        log.info("  Raw item keys: %s", list(raw_items[0].keys()))
        log.info("  Sample linkedinUrls (first 3): %s",
                 [i.get("linkedinUrl", "NO_linkedinUrl") for i in raw_items[:3]])

    items: list[dict] = []
    only_easy_apply = _cfg.get("only_easy_apply", False)

    for item in raw_items:
        url         = (_safe_str(item.get("linkedinUrl")) or _safe_str(item.get("easyApplyUrl"))).strip()
        description = (_safe_str(item.get("descriptionText")) or _safe_str(item.get("descriptionHtml"))).strip()
        loc         = _safe_str(item.get("location")).strip() or location
        has_easy_apply = bool(_safe_str(item.get("easyApplyUrl")).strip())

        # Filter by Easy Apply if enabled
        if only_easy_apply and not has_easy_apply:
            continue

        items.append({
            "title":       _safe_str(item.get("title")).strip(),
            "company":     _extract_company(item.get("company")).strip(),
            "location":    loc,
            "url":         url,
            "description": description,
            "posted_date": _safe_str(item.get("postedDate") or item.get("posted_date")),
            "source":      "LinkedIn",
            "has_easy_apply": has_easy_apply,
        })

    log.info("  → %d mapped items; jobs with no URL: %d",
             len(items), sum(1 for j in items if not j.get("url")))
    if only_easy_apply:
        log.info("  Easy Apply filter ON: %d items with Easy Apply", len(items))
    log.info("  Sample mapped URLs (first 3): %s",
             [j.get("url", "NO_URL") for j in items[:3]])
    if raw_items:
        log.info("  LinkedIn query field from result: %s", raw_items[0].get("query"))
    return items


def _deduplicate(jobs: list[dict]) -> list[dict]:
    seen_keys: set[str] = set()
    unique: list[dict] = []
    for job in jobs:
        url     = (job.get("url") or "").strip()
        title   = (job.get("title") or "").strip().lower()
        company = _extract_company(job.get("company")).strip().lower()
        # Prefer URL as key; fall back to company|title for URL-less jobs
        key = url if url else (f"{company}|{title}" if (company or title) else "")
        if not key or key in seen_keys:
            continue
        seen_keys.add(key)
        unique.append(job)
    log.info("Deduplication: %d → %d unique jobs", len(jobs), len(unique))
    return unique


def run(cfg: dict | None = None) -> list[dict]:
    reset_stop_flag()  # Reset stop flag at start of new session

    if cfg is None:
        cfg = load_config()

    roles = [r for r in (cfg.get("roles") or []) if isinstance(r, str) and r.strip()]
    if not roles:
        raise ValueError("cfg['roles'] is empty — select at least one role in the dashboard before scraping")

    smart_scrape = cfg.get("smart_scrape", True)

    # Load the persistent cache before clearing the run tables
    from dedup import db as _db_cache
    _db_cache.init_db()
    cached_urls = _db_cache.get_cached_job_urls()
    applied_urls, applied_combos = _get_applied_data()
    log.info("Cache: %d known URLs, %d already applied", len(cached_urls), len(applied_urls))

    # Clear previous run's temp data before fetching fresh results
    _clear_run_tables()

    client = _client()
    all_jobs: list[dict] = []

    # ── Stage 1: Broad scrape ────────────────────────────────────────────────
    combos = list(product(roles, cfg["locations"]))
    _scrape_progress.update({
        "total_combinations":     len(combos),
        "completed_combinations": 0,
        "current_search":         "",
        "jobs_found_so_far":      0,
    })

    _apify_success_count = 0
    _apify_error_count   = 0

    for role, location in combos:
        # Check stop flag before processing each combo
        if _should_stop_scraping():
            log.info("⏹️ Scraping stopped by user after %d combinations", _scrape_progress["completed_combinations"])
            break

        _scrape_progress["current_search"] = f"{role} @ {location}"
        try:
            batch = _scrape_linkedin(client, role, location, cfg)
            log.debug("Batch %s @ %s — %d jobs", role, location, len(batch))
            all_jobs.extend(batch)
            _apify_success_count += 1
        except RuntimeError as exc:
            # Catch the "stopped by user" exception
            if "stopped by user" in str(exc):
                log.info("⏹️ Scraper stopped: %s", exc)
                break
            log.warning("LinkedIn failed for %s @ %s: %s", role, location, exc)
            _apify_error_count += 1
        except Exception as exc:
            log.warning("LinkedIn failed for %s @ %s: %s", role, location, exc)
            _apify_error_count += 1
        finally:
            _scrape_progress["completed_combinations"] += 1
            _scrape_progress["jobs_found_so_far"] = len(all_jobs)

    if _apify_error_count and not _apify_success_count:
        log.warning(
            "⚠️ All %d Apify calls failed — will fall back to cached jobs. "
            "Check APIFY_API_TOKEN and actor ID in .env.",
            _apify_error_count,
        )
    elif _apify_error_count:
        log.warning("%d/%d Apify calls failed", _apify_error_count,
                    _apify_success_count + _apify_error_count)

    log.info("Stage 1 — scraped: %d jobs across all role/location combos", len(all_jobs))
    jobs = _deduplicate(all_jobs)
    jobs = [j for j in jobs if j.get("description")]

    # Upsert all fresh unique jobs into the persistent seen_jobs cache
    _db_cache.upsert_seen_jobs(jobs)

    # Backfill resume_hash for any seen_jobs rows missing it so the matcher
    # cache lookup never misses a previously-scored job due to a NULL hash
    try:
        current_resume_hash = _db_cache.get_resume_hash(cfg)
        if current_resume_hash:
            import sqlite3 as _sq
            _dbp = Path(__file__).parent.parent / "uploads" / "jobhunt.db"
            _dbc = _sq.connect(str(_dbp), timeout=10)
            _dbc.execute("PRAGMA journal_mode=WAL")
            _dbc.execute(
                "UPDATE seen_jobs SET resume_hash = ? WHERE resume_hash IS NULL AND dismissed = 0",
                (current_resume_hash,),
            )
            _dbc.commit()
            _dbc.close()
    except Exception as _rhe:
        log.warning("resume_hash backfill failed: %s", _rhe)

    # ── Stage 2: Title relevance filter (free, no API) ───────────────────────
    scraped_total = len(jobs)
    if smart_scrape:
        role_set = roles
        relevant = []
        for job in jobs:
            title = job.get("title", "")
            if any(is_title_relevant(title, r) for r in role_set):
                relevant.append(job)
            else:
                log.debug("Title filter drop: %s @ %s", title, job.get("company", ""))
        saved_calls = scraped_total - len(relevant)
        log.info(
            "Smart scrape: %d scraped → %d passed title filter → "
            "ready for Claude matching (saved %d API calls)",
            scraped_total, len(relevant), saved_calls,
        )
        jobs = relevant
    else:
        jobs = jobs[:_MAX_TOTAL_JOBS]

    # ── Applied dedup (URL + company+title) ─────────────────────────────────
    before = len(jobs)
    filtered: list[dict] = []
    for j in jobs:
        url     = j.get("url") or ""
        company = (j.get("company") or "").lower().strip()
        title   = (j.get("title") or "").lower().strip()
        combo   = (company, title)
        if url and url in applied_urls:
            log.info("Skip (already in dashboard): %s @ %s", j.get("title"), j.get("company"))
            continue
        if company and title and combo in applied_combos:
            log.info("Skip (already in dashboard): %s @ %s", j.get("title"), j.get("company"))
            continue
        filtered.append(j)
    skipped = before - len(filtered)
    if skipped:
        log.info("Skipped %d jobs already in applications table", skipped)
    jobs = filtered

    # ── Dismissed dedup (URL + company+title) ───────────────────────────────────
    try:
        with _db_cache._conn() as _dc:
            _drows = _dc.execute(
                "SELECT url, company, title FROM seen_jobs WHERE dismissed=1"
            ).fetchall()
        _dismissed_urls   = {r[0].strip() for r in _drows if r[0] and r[0].strip()}
        _dismissed_combos = {
            (r[1].lower().strip(), r[2].lower().strip())
            for r in _drows if r[1] and r[2]
        }
    except Exception:
        _dismissed_urls, _dismissed_combos = set(), set()
    if _dismissed_urls or _dismissed_combos:
        _before_d = len(jobs)
        _filtered_d = []
        for j in jobs:
            _url = j.get("url") or ""
            _co  = (j.get("company") or "").lower().strip()
            _ti  = (j.get("title")   or "").lower().strip()
            if _url and _url in _dismissed_urls:
                log.info("Skip (dismissed): %s @ %s", j.get("title"), j.get("company"))
                continue
            if _co and _ti and (_co, _ti) in _dismissed_combos:
                log.info("Skip (dismissed): %s @ %s", j.get("title"), j.get("company"))
                continue
            _filtered_d.append(j)
        _skipped_d = _before_d - len(_filtered_d)
        if _skipped_d:
            log.info("Skipped %d dismissed jobs", _skipped_d)
        jobs = _filtered_d

    # Tag fresh jobs: "new" if not seen before, "cached" if already in the cache
    for j in jobs:
        j["cache_status"] = "cached" if j.get("url") in cached_urls else "new"

    # ── Stage 3: Add cached jobs within the configured time window ─────────────
    only_easy_apply = cfg.get("only_easy_apply", False)
    fresh_urls    = {j.get("url") for j in jobs if j.get("url")}
    all_cached    = _db_cache.get_relevant_cached_jobs(cfg)
    cached_extra = []
    for j in all_cached:
        url     = j.get("url") or ""
        company = (j.get("company") or "").lower().strip()
        title   = (j.get("title") or "").lower().strip()
        if url and url in fresh_urls:
            continue
        if url and url in applied_urls:
            continue
        if company and title and (company, title) in applied_combos:
            continue
        if only_easy_apply and not j.get("has_easy_apply"):
            continue
        cached_extra.append(j)
    for j in cached_extra:
        j["cache_status"] = "cached"

    # Count jobs sitting in the cache but outside the time window (for logging)
    win_stats = _db_cache.get_cache_window_stats(cfg)
    outside_window = win_stats.get("outside_window", 0)

    jobs = jobs + cached_extra
    jobs = jobs[:_MAX_TOTAL_JOBS]

    new_count    = sum(1 for j in jobs if j.get("cache_status") == "new")
    cached_count = sum(1 for j in jobs if j.get("cache_status") == "cached")
    limit_label  = cfg.get("posted_limit", "24h")
    log.info(
        "Cache: %d jobs loaded from cache (posted within last %s) | "
        "%d new jobs scraped from LinkedIn | "
        "%d jobs in cache but outside time window (ignored) | "
        "%d total to process",
        cached_count, limit_label, new_count, outside_window, len(jobs),
    )

    # ── Filter cached jobs if skip_cached_jobs is enabled ───────────────────────
    # Only honour the toggle when the Apify scrape actually returned fresh jobs.
    # If the scrape produced nothing (API failure, rate-limit, etc.) we always
    # fall back to the cache so the pipeline doesn't silently return 0 jobs.
    if cfg.get("skip_cached_jobs", False):
        if new_count == 0:
            log.info(
                "skip_cached_jobs is ON but scrape returned 0 fresh jobs — "
                "keeping %d cached jobs as fallback",
                cached_count,
            )
        else:
            original_count = len(jobs)
            jobs = [j for j in jobs if j.get("cache_status") != "cached"]
            skipped = original_count - len(jobs)
            if skipped > 0:
                log.info("Skipped %d cached jobs (skip_cached_jobs enabled)", skipped)

    _save_to_db(jobs)
    return jobs


def load_all_cached(cfg: dict | None = None) -> list[dict]:
    """Populate scraped_jobs from every saved job in the persistent cache,
    ignoring the posted-date time window and without calling Apify at all.

    For jobs that were scraped in the past but never matched (e.g. they fell
    outside the posted_limit window before matching ran), this is the only way
    to see them again without a fresh scrape. Deliberately ignores
    only_easy_apply — this is a review/recovery view of everything saved, not
    a new scrape, and most historical rows predate has_easy_apply being tracked
    at all so filtering here would just hide almost everything.
    """
    if cfg is None:
        cfg = load_config()

    from dedup import db as _db_cache
    _db_cache.init_db()

    jobs = _db_cache.get_all_cached_jobs()
    for j in jobs:
        j["cache_status"] = "cached"

    _clear_run_tables()
    _save_to_db(jobs)
    log.info("Loaded %d saved jobs from cache (no scrape performed)", len(jobs))
    return jobs


def load_scored_cached(cfg: dict | None = None) -> list[dict]:
    """Populate scraped_jobs + matched_jobs from every job that already has a
    match_score from a past run, without re-running the matcher or Apify.

    The matcher only scores jobs it hasn't seen before, so a job matched in an
    earlier run (before matched_jobs got cleared by a later scrape) would
    otherwise be invisible in Step 3 unless it happens to also be unscored
    again. This recovers the full scored history from the persistent cache.
    """
    if cfg is None:
        cfg = load_config()

    from dedup import db as _db_cache
    _db_cache.init_db()

    jobs = _db_cache.get_all_scored_jobs()
    for j in jobs:
        j["cache_status"] = "cached"

    _clear_run_tables()
    _save_to_db(jobs)
    populated = _db_cache.bulk_populate_matched_jobs(jobs)
    log.info("Loaded %d previously-scored jobs from cache (%d matched_jobs rows created)",
              len(jobs), populated)
    return jobs


_DB_PATH = Path(__file__).parent.parent / "uploads" / "jobhunt.db"


def _db_connect() -> sqlite3.Connection:
    _DB_PATH.parent.mkdir(exist_ok=True)
    con = sqlite3.connect(str(_DB_PATH), timeout=15)
    con.execute("PRAGMA journal_mode=WAL")
    return con


def _clear_run_tables() -> None:
    """Delete all rows from the two temp tables before a new scrape run."""
    try:
        con = _db_connect()
        try:
            con.execute("DELETE FROM matched_jobs")
            con.execute("DELETE FROM scraped_jobs")
            con.commit()
            log.info("Cleared scraped_jobs and matched_jobs for fresh run")
        finally:
            con.close()
    except Exception as exc:
        log.warning("Could not clear run tables: %s", exc)


def _get_applied_data() -> tuple[set[str], set[tuple[str, str]]]:
    """Return (applied_urls, applied_combos) from the applications table."""
    try:
        con = _db_connect()
        try:
            rows = con.execute(
                "SELECT job_url, company, role FROM applications"
                " WHERE job_url IS NOT NULL OR company IS NOT NULL"
            ).fetchall()
            applied_urls   = {r[0].strip() for r in rows if r[0] and r[0].strip()}
            applied_combos = {
                (r[1].lower().strip(), r[2].lower().strip())
                for r in rows if r[1] and r[2]
            }
            log.info(
                "Applied dedup: %d URLs, %d company+title combos",
                len(applied_urls), len(applied_combos),
            )
            return applied_urls, applied_combos
        finally:
            con.close()
    except Exception as exc:
        log.warning("Could not fetch applied data from DB: %s", exc)
        return set(), set()


def _save_to_db(jobs: list[dict]) -> None:
    """Insert fresh scraped jobs into SQLite (tables already cleared by _clear_run_tables)."""
    _DB_PATH.parent.mkdir(exist_ok=True)
    con = _db_connect()
    try:
        con.executemany(
            "INSERT OR IGNORE INTO scraped_jobs"
            " (title, company, location, url, description, source, posted_date, has_easy_apply, cache_status)"
            " VALUES (?,?,?,?,?,?,?,?,?)",
            [
                (
                    j.get("title", ""),
                    _extract_company(j.get("company")),
                    j.get("location", ""),
                    j.get("url", ""),
                    j.get("description", ""),
                    j.get("source", "LinkedIn"),
                    j.get("posted_date", ""),
                    1 if j.get("has_easy_apply") else 0,
                    j.get("cache_status", "new"),
                )
                for j in jobs
            ],
        )
        con.commit()
        log.info("Saved %d scraped jobs to DB", len(jobs))
    except Exception as exc:
        log.error("Failed to save scraped jobs to DB: %s", exc)
        con.rollback()
    finally:
        con.close()


def extract_job_from_url(url: str, max_attempts: int = 3) -> dict | None:
    """Fetch a job URL and extract title, company, location, description.

    Uses Playwright to load the page and JavaScript to extract common job fields.
    Job sites occasionally serve a slow or empty page on a single attempt
    (rate limiting, a transient block, a cold connection) — retry a few times
    with a short backoff before giving up.
    Returns a job dict ready for DB insertion, or None if extraction fails.
    """
    import asyncio
    from playwright.async_api import async_playwright

    async def _fetch_once():
        browser = None
        try:
            async with async_playwright() as p:
                # Playwright's own bundled Chromium isn't installed on this
                # machine — the rest of the codebase always launches the
                # user's real installed Chrome via channel="chrome" instead.
                browser = await p.chromium.launch(channel="chrome")
                page = await browser.new_page()
                # networkidle times out on most real job sites (analytics/chat
                # widgets keep the network busy indefinitely) — domcontentloaded
                # plus a short settle delay is far more reliable here.
                await page.goto(url, wait_until="domcontentloaded", timeout=20000)
                await page.wait_for_timeout(1500)

                data = await page.evaluate("""() => {
                    const getText = (sel) => {
                        const el = document.querySelector(sel);
                        return el ? el.innerText.trim() : '';
                    };

                    const title = getText('h1') || getText('h2') || getText('[data-job-title]') || getText('.job-title');
                    const company = getText('.topcard__org-name-link') || getText('[data-company]') || getText('.company-name') || getText('.employer') || getText('[data-company-name]');
                    const location = getText('.topcard__flavor--bullet') || getText('[data-location]') || getText('.location') || getText('.job-location');
                    let description = getText('.job-description') || getText('[data-description]') || getText('.description') || getText('#job-description');

                    if (!description) {
                        const main = document.querySelector('[role="main"]');
                        description = main ? main.innerText : document.body.innerText;
                    }

                    return {
                        title: title.slice(0, 300),
                        company: company.slice(0, 200),
                        location: location.slice(0, 200),
                        description: description.slice(0, 3000)
                    };
                }""")

                if not data.get("title"):
                    raise ValueError("no title found on page")

                return {
                    "url": url,
                    "title": data.get("title", ""),
                    "company": data.get("company", "") or "Unknown",
                    "location": data.get("location", ""),
                    "description": data.get("description", ""),
                    "source": "Manual",
                    "posted_date": "",
                }
        finally:
            if browser:
                try:
                    await browser.close()
                except Exception:
                    pass

    async def _fetch_with_retries():
        for attempt in range(1, max_attempts + 1):
            try:
                return await _fetch_once()
            except Exception as e:
                log.warning(f"Extraction attempt {attempt}/{max_attempts} failed for {url}: {e}")
                if attempt < max_attempts:
                    await asyncio.sleep(2)
        log.warning(f"Giving up on {url} after {max_attempts} attempts")
        return None

    try:
        return asyncio.run(_fetch_with_retries())
    except Exception as e:
        log.error(f"Error extracting job from {url}: {e}")
        return None


if __name__ == "__main__":
    import json as _json
    from core.config import setup_logging
    cfg = load_config()
    setup_logging(cfg)
    results = run(cfg)
    print(_json.dumps(results[:3], indent=2, ensure_ascii=False))
    print(f"\nTotal: {len(results)} jobs")