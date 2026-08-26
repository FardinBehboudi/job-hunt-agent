"""
tailor.py — generate a tailored resume + cover letter for each matched job.

Pipeline:
  1. Claude suggests changes (JSON) → applied with python-docx
  2. .docx → PDF via mammoth (docx→HTML) + weasyprint (HTML→PDF)
     Falls back to LibreOffice soffice if weasyprint is unavailable.

Base resume must be uploads/resume_en.docx (cfg["paths"]["resume_en_docx"]).
Cover letter template can be .docx or .pdf (text extracted via pdfplumber).
"""

import json
import logging
import os
import re
import shutil
from datetime import date
from pathlib import Path

import anthropic
import pdfplumber
from docx import Document
from dotenv import load_dotenv

from core.config import load_config

load_dotenv()
log = logging.getLogger(__name__)

_MODEL = "claude-sonnet-4-6"
_WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _slugify(text: str) -> str:
    return re.sub(r"[^\w-]", "_", text.strip())[:40]


def _archive_folder(job: dict, cfg: dict) -> Path:
    history: Path = cfg["paths"]["history_folder"]
    folder = history / f"{_slugify(job.get('company', 'Unknown'))}_{_slugify(job.get('title', 'Role'))}_{date.today().isoformat()}"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def _company_tailored_folder(job: dict, cfg: dict) -> Path:
    """The per-company folder the applier reads tailored resumes/cover letters
    from — one folder per company, refreshed each time tailoring runs for it."""
    base: Path = cfg["paths"]["resume_tailored_dir"]
    folder = base / _slugify(job.get("company", "Unknown"))
    folder.mkdir(parents=True, exist_ok=True)
    return folder


# ── PDF export ────────────────────────────────────────────────────────────────

def _find_soffice() -> str:
    """Locate the LibreOffice binary. Not always on PATH on Windows installs."""
    import shutil as _shutil
    on_path = _shutil.which("soffice") or _shutil.which("soffice.exe")
    if on_path:
        return on_path
    for candidate in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/opt/libreoffice/program/soffice",
    ):
        if Path(candidate).exists():
            return candidate
    return "soffice"  # last resort — let subprocess raise FileNotFoundError


def _export_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert .docx → PDF.  Primary: mammoth + weasyprint.  Fallback: LibreOffice."""
    try:
        import mammoth
        import weasyprint

        with open(docx_path, "rb") as f:
            result = mammoth.convert_to_html(f)
        html_body = result.value

        html = (
            "<!DOCTYPE html><html lang='en'><head>"
            "<meta charset='utf-8'>"
            "<style>"
            "@page{margin:2cm}"
            "body{font-family:Arial,sans-serif;font-size:10pt;line-height:1.4}"
            "h1{font-size:16pt;margin-bottom:4pt}"
            "h2{font-size:12pt;margin-top:12pt;margin-bottom:4pt;border-bottom:1px solid #ccc}"
            "p,li{margin:4pt 0}"
            "ul{padding-left:20pt}"
            "</style></head>"
            f"<body>{html_body}</body></html>"
        )
        weasyprint.HTML(string=html).write_pdf(str(pdf_path))
        log.info("PDF exported (weasyprint): %s", pdf_path.name)
        return

    except ImportError as exc:
        log.warning("mammoth/weasyprint not available (%s) — trying LibreOffice", exc)
    except Exception as exc:
        log.warning("weasyprint export failed (%s) — trying LibreOffice", exc)

    # Fallback: LibreOffice headless
    try:
        import subprocess
        result = subprocess.run(
            [_find_soffice(), "--headless", "--convert-to", "pdf",
             "--outdir", str(pdf_path.parent), str(docx_path)],
            capture_output=True, timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr.decode())
        generated = pdf_path.parent / (docx_path.stem + ".pdf")
        if generated.exists() and generated != pdf_path:
            generated.rename(pdf_path)
        log.info("PDF exported (LibreOffice): %s", pdf_path.name)
        return
    except FileNotFoundError:
        pass
    except Exception as exc:
        log.warning("LibreOffice export also failed (%s)", exc)

    # Last resort: keep the .docx for manual conversion
    docx_dest = pdf_path.with_suffix(".docx")
    try:
        if docx_dest != docx_path:
            shutil.copy2(docx_path, docx_dest)
        log.warning("No PDF converter available — using .docx instead: %s", docx_dest)
    except Exception as exc:
        log.warning("Could not save fallback .docx (%s) — resume/CL still available at %s", exc, docx_path)


# ── Section reordering ────────────────────────────────────────────────────────

def _reorder_sections(doc: Document, section_order: list[str]) -> None:
    """Reorder document body sections in-place according to section_order."""
    if not section_order:
        return

    body = doc.element.body

    def get_text(el) -> str:
        return "".join(t.text or "" for t in el.iter(f"{{{_WML}}}t"))

    def is_heading(el) -> bool:
        pStyle = el.find(f".//{{{_WML}}}pStyle")
        if pStyle is None:
            return False
        val = pStyle.get(f"{{{_WML}}}val", "")
        return "heading" in val.lower()

    # sectPr must stay last — pull it out
    sect_pr = body.find(f"{{{_WML}}}sectPr")
    all_children = [c for c in list(body) if c is not sect_pr]

    # Group children: header_block (pre-first-heading) + sections list
    header_block: list = []
    sections: list[tuple[str, object, list]] = []  # (heading_text, heading_el, content_els)
    current: list | None = None

    for child in all_children:
        if child.tag == f"{{{_WML}}}p" and is_heading(child):
            heading_text = get_text(child).strip()
            current = []
            sections.append((heading_text, child, current))
        elif current is None:
            header_block.append(child)
        else:
            current.append(child)

    # Match requested names to actual sections (case-insensitive substring)
    remaining = list(sections)
    ordered: list[tuple] = []
    for req in section_order:
        for sec in remaining:
            if req.lower() in sec[0].lower() or sec[0].lower() in req.lower():
                ordered.append(sec)
                remaining.remove(sec)
                break
    ordered.extend(remaining)  # append unmatched sections at end

    # Rebuild body
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    for el in header_block:
        body.append(el)
    for _, h_el, content in ordered:
        body.append(h_el)
        for el in content:
            body.append(el)
    if sect_pr is not None:
        body.append(sect_pr)


# ── Resume tailoring ──────────────────────────────────────────────────────────

_RESUME_SYSTEM = """\
You are a professional resume coach helping a candidate tailor their resume.
IMPORTANT: Never invent experience or skills. Only reorder, re-emphasise, or surface
keywords that are genuinely present in the candidate's background.
Return ONLY valid JSON — no markdown fences, no prose.
"""

_RESUME_PROMPT = """\
## Job Description
{jd}

## Current Resume Text
{resume}

Return a JSON object:
{{
  "summary_tweak": "<revised 2-3 sentence professional summary, or null>",
  "keywords_to_add": ["keyword1", "keyword2"],
  "section_order": ["Experience", "Skills", "Education"],
  "bullet_replacements": [
    {{"original": "exact bullet text from resume", "replacement": "improved version"}}
  ]
}}

All changes must be factually accurate to the resume provided.
Only include bullet_replacements where you can quote the original text exactly.
"""

_CL_SYSTEM = """\
You are a professional cover letter writer.
Preserve the candidate's voice and style from the template.
Return ONLY the full cover letter body text — no JSON, no metadata, no subject line.
"""

_CL_PROMPT = """\
## Cover Letter Template
{template}

## Job Description
{jd}

## Company Name
{company}

Write a tailored cover letter for this specific role and company.
Preserve the original tone, structure, and personal voice exactly.
Do not fabricate achievements — draw only from what the template shows.
"""

# ── German versions ─────────────────────────────────────────────────────────
# The surgical bullet-replacement approach used for the English resume can't
# reliably produce a full translation, so the German resume is generated as a
# complete rewrite instead and rendered into a fresh .docx (see
# _build_resume_docx_from_text below).

_RESUME_DE_SYSTEM = """\
You are a professional resume writer producing a German-language resume (Lebenslauf).
Translate and tailor the candidate's resume into professional German for this job.
IMPORTANT: Never invent experience or skills — translate and re-emphasise only what
is genuinely present in the original resume.
Return ONLY the full resume text in German, as blank-line-separated blocks. Each
section starts with a short German heading on its own line (e.g. "Profil",
"Berufserfahrung", "Fähigkeiten", "Ausbildung"), followed by its content. Use
"- " to prefix bullet points.
"""

_RESUME_DE_PROMPT = """\
## Job Description
{jd}

## Original Resume (English)
{resume}

Produce a complete, professional German-language resume tailored to this job.
Preserve the structure: contact info, professional summary, work experience with
bullet points, skills, education. Do not fabricate anything not present in the
original.
"""

_CL_DE_SYSTEM = """\
You are a professional cover letter writer producing a German-language cover
letter (Anschreiben). Preserve the candidate's voice and style from the
template, translated and adapted naturally to German business-letter conventions.
Return ONLY the full cover letter body text in German — no JSON, no metadata,
no subject line.
"""

_CL_DE_PROMPT = """\
## Cover Letter Template (English)
{template}

## Job Description
{jd}

## Company Name
{company}

Write a tailored German-language cover letter (Anschreiben) for this specific
role and company. Preserve the original tone and structure, adapted naturally
to German conventions. Do not fabricate achievements — draw only from what the
template shows.
"""


def _read_docx_text(path: Path) -> str:
    """Extract all text from a .docx — body paragraphs plus table cells.
    Many resumes lay their actual content out in tables, which doc.paragraphs
    alone skips entirely."""
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _read_pdf_text(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n".join(parts)


def _read_template_text(path: Path) -> str:
    return _read_pdf_text(path) if path.suffix.lower() == ".pdf" else _read_docx_text(path)


def _apply_resume_changes(base_docx: Path, out_docx: Path, changes: dict) -> None:
    doc = Document(base_docx)

    # 1. Reorder sections
    section_order = changes.get("section_order", [])
    if section_order:
        _reorder_sections(doc, section_order)

    # 2. Replace professional summary
    summary = changes.get("summary_tweak")
    if summary:
        for para in doc.paragraphs:
            if len(para.text) > 60 and para.style.name.startswith("Normal"):
                para.clear()
                para.add_run(summary)
                break

    # 3. Apply bullet replacements (match against exact original text)
    replacements = {r["original"]: r["replacement"]
                    for r in changes.get("bullet_replacements", [])
                    if r.get("original") and r.get("replacement")}
    for para in doc.paragraphs:
        for orig, repl in replacements.items():
            if orig in para.text:
                for run in para.runs:
                    if orig in run.text:
                        run.text = run.text.replace(orig, repl)
                        break

    # 4. Add keywords to Skills section
    keywords = changes.get("keywords_to_add", [])
    if keywords:
        in_skills = False
        for para in doc.paragraphs:
            text_lower = para.text.strip().lower()
            if "skill" in text_lower and para.style.name.lower().startswith("heading"):
                in_skills = True
                continue
            if in_skills and para.style.name.lower().startswith("heading"):
                break
            if in_skills and para.text.strip() and para.runs:
                new_kw = [k for k in keywords if k.lower() not in para.text.lower()]
                if new_kw:
                    para.runs[-1].text += ", " + ", ".join(new_kw)
                break

    doc.save(out_docx)


def _apply_cover_letter(template_docx: Path, out_docx: Path, body_text: str) -> None:
    doc = Document(template_docx)
    body_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    # Preserve first 2 paragraphs (name/address header), clear the rest
    for para in body_paragraphs[2:]:
        para.clear()
    # Append new body text
    for line in body_text.split("\n\n"):
        line = line.strip()
        if line:
            p = doc.add_paragraph(line)
            p.style = doc.styles["Normal"]
    doc.save(out_docx)


def _build_cl_docx_from_text(body_text: str, out_path: Path) -> None:
    """Build a cover letter .docx from scratch (used when template is a PDF)."""
    doc = Document()
    for para_text in body_text.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)
    doc.save(out_path)


def _build_resume_docx_from_text(body_text: str, out_path: Path) -> None:
    """Build a resume .docx from scratch from a full Claude-generated rewrite
    (used for the German version — see _RESUME_DE_PROMPT). Blank-line-separated
    blocks; a block's short, punctuation-free first line becomes a heading,
    lines starting with '- '/'• '/'* ' become bullet points."""
    doc = Document()
    for block in body_text.split("\n\n"):
        lines = [l.strip() for l in block.strip().split("\n") if l.strip()]
        if not lines:
            continue
        first = lines[0]
        is_heading = len(first) < 50 and not first.endswith((".", ",", ":")) and len(lines) > 1
        body_lines = lines[1:] if is_heading else lines
        if is_heading:
            doc.add_heading(first, level=2)
        for line in body_lines:
            if line.startswith(("- ", "• ", "* ")):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)
    doc.save(out_path)


# ── Public API ────────────────────────────────────────────────────────────────

def create_docs(job: dict, cfg: dict | None = None) -> Path:
    """
    Build tailored resume + cover letter for *job*.
    Returns the archive folder path.
    """
    if cfg is None:
        cfg = load_config()

    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
    archive = _archive_folder(job, cfg)

    # Save raw job data
    (archive / "job_description.txt").write_text(job.get("description", ""), encoding="utf-8")
    (archive / "match_score.json").write_text(
        json.dumps({
            "match_score":           job.get("match_score"),
            "interview_chance":      job.get("interview_chance"),
            "german_level_required": job.get("german_level_required"),
            "match_summary":         job.get("match_summary"),
        }, indent=2),
        encoding="utf-8",
    )

    resume_docx  = cfg["paths"]["resume_en_docx"]
    cl_template  = cfg["paths"]["cover_letter_template"]
    company_dir  = _company_tailored_folder(job, cfg)

    if not resume_docx.exists():
        log.warning("resume_en.docx not found at %s — skipping resume tailoring", resume_docx)
    if not cl_template.exists():
        log.warning("cover_letter_template not found at %s — skipping CL", cl_template)

    def _mirror_to_company_dir(archive_file: Path, stem: str) -> None:
        """Copy a generated archive_file (resume_en.pdf, cover_letter_de.docx, ...)
        into the per-company folder the applier actually reads from."""
        if archive_file.exists():
            shutil.copy2(archive_file, company_dir / f"{stem}{archive_file.suffix}")

    # ── Tailor resume (EN) ────────────────────────────────────────────────────
    resume_text = ""
    if resume_docx.exists():
        resume_text = _read_docx_text(resume_docx)
        resp = client.messages.create(
            model=_MODEL,
            max_tokens=1024,
            system=_RESUME_SYSTEM,
            messages=[{"role": "user", "content": _RESUME_PROMPT.format(
                jd=job["description"][:5000],
                resume=resume_text[:4000],
            )}],
        )
        raw = resp.content[0].text.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```[a-z]*\n?", "", raw)
            raw = re.sub(r"\n?```$", "", raw)
            raw = raw.strip()
        try:
            changes = json.loads(raw)
        except json.JSONDecodeError:
            log.warning("Resume tailoring returned invalid JSON — using base resume unchanged")
            changes = {}

        out_resume_docx = archive / "resume_en.docx"
        _apply_resume_changes(resume_docx, out_resume_docx, changes)
        _export_docx_to_pdf(out_resume_docx, archive / "resume_en.pdf")
        log.info("Resume (EN) tailored → %s", archive.name)

        _mirror_to_company_dir(out_resume_docx, "resume_en")
        _mirror_to_company_dir(archive / "resume_en.pdf", "resume_en")

    # ── Tailor resume (DE) — full rewrite, not a surgical edit ──────────────────
    if resume_docx.exists() and resume_text:
        resp = client.messages.create(
            model=_MODEL,
            max_tokens=1600,
            system=_RESUME_DE_SYSTEM,
            messages=[{"role": "user", "content": _RESUME_DE_PROMPT.format(
                jd=job["description"][:5000],
                resume=resume_text[:4000],
            )}],
        )
        resume_de_body = resp.content[0].text.strip()

        out_resume_de_docx = archive / "resume_de.docx"
        _build_resume_docx_from_text(resume_de_body, out_resume_de_docx)
        _export_docx_to_pdf(out_resume_de_docx, archive / "resume_de.pdf")
        log.info("Resume (DE) tailored → %s", archive.name)

        _mirror_to_company_dir(out_resume_de_docx, "resume_de")
        _mirror_to_company_dir(archive / "resume_de.pdf", "resume_de")

    # ── Tailor cover letter (EN) ──────────────────────────────────────────────
    cl_text = ""
    if cl_template.exists():
        cl_text = _read_template_text(cl_template)
        resp = client.messages.create(
            model=_MODEL,
            max_tokens=1024,
            system=_CL_SYSTEM,
            messages=[{"role": "user", "content": _CL_PROMPT.format(
                template=cl_text,
                jd=job["description"][:4000],
                company=job.get("company", "the company"),
            )}],
        )
        cl_body = resp.content[0].text.strip()
        out_cl_docx = archive / "cover_letter_en.docx"
        if cl_template.suffix.lower() == ".docx":
            _apply_cover_letter(cl_template, out_cl_docx, cl_body)
        else:
            _build_cl_docx_from_text(cl_body, out_cl_docx)
        _export_docx_to_pdf(out_cl_docx, archive / "cover_letter_en.pdf")
        log.info("Cover letter (EN) tailored → %s", archive.name)

        _mirror_to_company_dir(out_cl_docx, "cover_letter_en")
        _mirror_to_company_dir(archive / "cover_letter_en.pdf", "cover_letter_en")

    # ── Tailor cover letter (DE) ──────────────────────────────────────────────
    if cl_template.exists() and cl_text:
        resp = client.messages.create(
            model=_MODEL,
            max_tokens=1024,
            system=_CL_DE_SYSTEM,
            messages=[{"role": "user", "content": _CL_DE_PROMPT.format(
                template=cl_text,
                jd=job["description"][:4000],
                company=job.get("company", "the company"),
            )}],
        )
        cl_de_body = resp.content[0].text.strip()
        out_cl_de_docx = archive / "cover_letter_de.docx"
        _build_cl_docx_from_text(cl_de_body, out_cl_de_docx)
        _export_docx_to_pdf(out_cl_de_docx, archive / "cover_letter_de.pdf")
        log.info("Cover letter (DE) tailored → %s", archive.name)

        _mirror_to_company_dir(out_cl_de_docx, "cover_letter_de")
        _mirror_to_company_dir(archive / "cover_letter_de.pdf", "cover_letter_de")

    log.info("Archive: %s · company folder: %s", archive, company_dir)
    return archive


if __name__ == "__main__":
    from core.config import setup_logging
    cfg = load_config()
    setup_logging(cfg)
    sample_job = {
        "title": "Data Engineer",
        "company": "TestCorp GmbH",
        "location": "Berlin",
        "url": "https://example.com/job/1",
        "description": "Looking for a Data Engineer with Python, Airflow, dbt, and SQL.",
        "source": "LinkedIn",
        "match_score": 85,
        "interview_chance": "high",
        "german_level_required": "B1",
        "match_summary": "Strong match on Python and data pipeline skills.",
    }
    path = create_docs(sample_job, cfg)
    print(f"Archive: {path}")
